from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def clean_inline_md(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    text = re.sub(r"\[(\d+)\]\(([^\)]+)\)", r"\1", text)
    return text.strip()


def is_table_separator(line: str) -> bool:
    core = line.strip().strip("|").strip()
    if not core:
        return False
    parts = [p.strip() for p in core.split("|")]
    return all(re.fullmatch(r":?-{3,}:?", p) is not None for p in parts)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        raw = line.strip()
        if not raw.startswith("|"):
            continue
        cols = [clean_inline_md(c.strip()) for c in raw.strip("|").split("|")]
        rows.append(cols)
    return rows


def apply_ieee_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(clean_inline_md(title))
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    apply_ieee_style(doc)

    i = 0
    title_added = False

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# ") and not title_added:
            add_title(doc, stripped[2:].strip())
            title_added = True
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(clean_inline_md(stripped[2:].strip()), level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(clean_inline_md(stripped[3:].strip()), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(clean_inline_md(stripped[4:].strip()), level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(clean_inline_md(stripped[5:].strip()), level=3)
            i += 1
            continue

        if stripped.startswith("|"):
            block: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1

            rows = parse_table_rows(block)
            if len(rows) >= 2 and is_table_separator(block[1]):
                header = rows[0]
                body = rows[2:]
                table = doc.add_table(rows=1 + len(body), cols=len(header))
                table.style = "Table Grid"

                for c, v in enumerate(header):
                    table.cell(0, c).text = v

                for r_idx, row in enumerate(body, start=1):
                    for c in range(len(header)):
                        table.cell(r_idx, c).text = row[c] if c < len(row) else ""
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph(clean_inline_md(" ".join(block)))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if re.match(r"^\d+\.\s+", stripped):
            item = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(clean_inline_md(item), style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(clean_inline_md(stripped[2:].strip()), style="List Bullet")
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith("#") or nxt.startswith("|"):
                break
            if re.match(r"^\d+\.\s+", nxt) or nxt.startswith("- "):
                break
            paragraph_lines.append(nxt)
            i += 1

        p = doc.add_paragraph(clean_inline_md(" ".join(paragraph_lines)))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(docx_path)


if __name__ == "__main__":
    src = Path("d:/tox-agent/docs/papers/ToxAgent_IEEE_ACM_Full_Paper_vi.md")
    out = Path("d:/tox-agent/docs/papers/ToxAgent_IEEE_ACM_Full_Paper_vi.docx")
    convert_markdown_to_docx(src, out)
    print(f"DOCX_WRITTEN={out}")
